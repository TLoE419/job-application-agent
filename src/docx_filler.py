"""
DOCX Template Filler
Replaces placeholders in DOCX templates with actual CV data from YAML
"""
from docx import Document
from pathlib import Path
import yaml
import re
from typing import Dict, Any
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


class DOCXFiller:
    """Fill DOCX templates with CV data"""

    @staticmethod
    def _add_hyperlink(paragraph, text, url, run_template):
        """
        Add a hyperlink to a paragraph

        Args:
            paragraph: The paragraph to add hyperlink to
            text: Display text for the hyperlink
            url: The URL to link to
            run_template: Template run to copy formatting from
        """
        # Ensure URL has protocol
        if not url.startswith('http://') and not url.startswith('https://'):
            url = 'https://' + url

        # Get paragraph's part
        part = paragraph.part
        r_id = part.relate_to(url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)

        # Create hyperlink element
        hyperlink = OxmlElement('w:hyperlink')
        hyperlink.set(qn('r:id'), r_id)

        # Create a new run for the hyperlink
        new_run = OxmlElement('w:r')
        rPr = OxmlElement('w:rPr')

        # Copy formatting from template run
        if run_template.font.name:
            rFonts = OxmlElement('w:rFonts')
            rFonts.set(qn('w:ascii'), run_template.font.name)
            rPr.append(rFonts)

        if run_template.font.size:
            sz = OxmlElement('w:sz')
            sz.set(qn('w:val'), str(run_template.font.size.pt * 2))
            rPr.append(sz)

        # Add blue color and underline for hyperlinks
        color = OxmlElement('w:color')
        color.set(qn('w:val'), '0563C1')  # Blue color for hyperlinks
        rPr.append(color)

        u = OxmlElement('w:u')
        u.set(qn('w:val'), 'single')
        rPr.append(u)

        new_run.append(rPr)

        # Add text
        t = OxmlElement('w:t')
        t.text = text
        new_run.append(t)

        hyperlink.append(new_run)
        paragraph._p.append(hyperlink)

    @staticmethod
    def fill_template(template_path: str, cv_data: Dict[str, Any], output_path: str) -> str:
        """
        Fill template with CV data

        Args:
            template_path: Path to template.docx file
            cv_data: CV data as dictionary (from YAML)
            output_path: Where to save the filled document

        Returns:
            Path to the generated document
        """
        # Load template
        doc = Document(template_path)

        # Build replacement mappings from CV data
        replacements = DOCXFiller._build_replacements(cv_data)

        print(f"📝 Filling template with {len(replacements)} placeholders...")

        # Replace in paragraphs (multiple passes to handle multiple placeholders per paragraph)
        for para in doc.paragraphs:
            # Keep replacing until no more placeholders found
            max_iterations = 20  # Prevent infinite loop
            for _ in range(max_iterations):
                full_text = para.text
                has_placeholder = any(placeholder in full_text for placeholder in replacements.keys())
                if not has_placeholder:
                    break
                DOCXFiller._replace_in_paragraph(para, replacements, cv_data)

        # Replace in tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        # Keep replacing until no more placeholders found
                        max_iterations = 20
                        for _ in range(max_iterations):
                            full_text = para.text
                            has_placeholder = any(placeholder in full_text for placeholder in replacements.keys())
                            if not has_placeholder:
                                break
                            DOCXFiller._replace_in_paragraph(para, replacements, cv_data)

        # Save filled document
        Path(output_path).parent.mkdir(parents=True, exist_ok=True)
        doc.save(output_path)

        print(f"✓ Document saved: {output_path}")
        return output_path

    @staticmethod
    def _build_replacements(cv_data: Dict[str, Any]) -> Dict[str, str]:
        """
        Build replacement dictionary from CV data

        Converts CV data structure to placeholder mappings:
        cv_data['personal_info']['name'] -> {{NAME}}
        cv_data['education'][0]['school'] -> {{EDU_1_SCHOOL}}
        """
        replacements = {}

        # Personal info / Header
        if 'personal_info' in cv_data:
            personal = cv_data['personal_info']
            # Support both {{NAME}} and {{Name}} variants
            replacements['{{NAME}}'] = personal.get('name', '')
            replacements['{{Name}}'] = personal.get('name', '')
            replacements['{{LOCATION}}'] = personal.get('location', '')
            replacements['{{EMAIL}}'] = personal.get('email', '')
            replacements['{{PHONE}}'] = personal.get('phone', '')
            replacements['{{GITHUB}}'] = personal.get('github', '')
            replacements['{{LINKEDIN}}'] = personal.get('linkedin', '')

        # Education
        if 'education' in cv_data:
            for idx, edu in enumerate(cv_data['education'], start=1):
                # Support both 'school' and 'institution'
                school = edu.get('school', edu.get('institution', ''))
                replacements[f'{{{{EDU_{idx}_SCHOOL}}}}'] = school

                replacements[f'{{{{EDU_{idx}_LOCATION}}}}'] = edu.get('location', '')
                replacements[f'{{{{EDU_{idx}_DEGREE}}}}'] = edu.get('degree', '')

                # Support both 'date' and 'start_date + end_date'
                date = edu.get('date', '')
                if not date and 'start_date' in edu and 'end_date' in edu:
                    date = f"{edu['start_date']} – {edu['end_date']}"
                replacements[f'{{{{EDU_{idx}_DATE}}}}'] = date

                # Handle courses (can be list or string)
                courses = edu.get('courses', '')
                if isinstance(courses, list):
                    courses = ', '.join(courses)
                if courses:
                    replacements[f'{{{{EDU_{idx}_COURSES}}}}'] = courses
                else:
                    replacements[f'{{{{EDU_{idx}_COURSES}}}}'] = ''

        # Experience
        if 'experience' in cv_data:
            for idx, exp in enumerate(cv_data['experience'], start=1):
                replacements[f'{{{{EXP_{idx}_COMPANY}}}}'] = exp.get('company', '')
                replacements[f'{{{{EXP_{idx}_LOCATION}}}}'] = exp.get('location', '')
                replacements[f'{{{{EXP_{idx}_TITLE}}}}'] = exp.get('title', '')

                # Support both 'date' and 'start_date + end_date'
                date = exp.get('date', '')
                if not date and 'start_date' in exp and 'end_date' in exp:
                    date = f"{exp['start_date']} – {exp['end_date']}"
                replacements[f'{{{{EXP_{idx}_DATE}}}}'] = date

                # Achievements
                achievements = exp.get('achievements', [])
                if isinstance(achievements, list):
                    for ach_idx, achievement in enumerate(achievements, start=1):
                        replacements[f'{{{{EXP_{idx}_ACHIEVEMENT_{ach_idx}}}}}'] = achievement

                    # Also create combined version with bullets
                    combined = '\n'.join(achievements)
                    replacements[f'{{{{EXP_{idx}_ACHIEVEMENTS}}}}'] = combined

        # Skills
        if 'skills' in cv_data:
            skills = cv_data['skills']

            # Handle different skill structures
            if isinstance(skills, dict):
                # Case 1: {languages: [...], frameworks_tools/frameworks_and_tools: [...]}
                if 'languages' in skills:
                    langs = skills['languages']
                    if isinstance(langs, list):
                        langs = ', '.join(langs)
                    replacements['{{SKILL_LANGUAGE}}'] = langs
                    replacements['{{SKILLS_LANGUAGES}}'] = langs

                # Support both frameworks_tools and frameworks_and_tools
                frameworks_key = 'frameworks_tools' if 'frameworks_tools' in skills else 'frameworks_and_tools'
                if frameworks_key in skills:
                    frameworks = skills[frameworks_key]
                    if isinstance(frameworks, list):
                        frameworks = ', '.join(frameworks)
                    replacements['{{SKILL_FRAME_TOOL}}'] = frameworks
                    replacements['{{SKILLS_FRAMEWORKS_TOOLS}}'] = frameworks

                # Case 2: {language_processing: "...", data_visualization: "..."}
                if 'language_processing' in skills:
                    replacements['{{SKILLS_LANGUAGE_PROCESSING}}'] = skills['language_processing']

                if 'data_visualization' in skills:
                    replacements['{{SKILLS_DATA_VIZ}}'] = skills['data_visualization']

        # Projects
        if 'projects' in cv_data:
            for idx, proj in enumerate(cv_data['projects'], start=1):
                replacements[f'{{{{PROJ_{idx}_NAME}}}}'] = proj.get('name', '')
                replacements[f'{{{{PROJ_{idx}_DATE}}}}'] = proj.get('date', '')

                # Descriptions (can be list or single string)
                descriptions = proj.get('descriptions', proj.get('description', []))
                if isinstance(descriptions, str):
                    descriptions = [descriptions]

                if isinstance(descriptions, list):
                    for desc_idx, desc in enumerate(descriptions, start=1):
                        replacements[f'{{{{PROJ_{idx}_DESC_{desc_idx}}}}}'] = desc

                    # Also create combined version
                    if len(descriptions) == 1:
                        replacements[f'{{{{PROJ_{idx}_DESC}}}}'] = descriptions[0]
                    else:
                        combined = '\n'.join(descriptions)
                        replacements[f'{{{{PROJ_{idx}_DESC}}}}'] = combined

        return replacements

    @staticmethod
    def _replace_in_paragraph(paragraph, replacements: Dict[str, str], cv_data: Dict[str, Any] = None):
        """
        Replace placeholders while preserving run-level formatting

        Strategy: Each run keeps its own formatting, we only replace text content
        For GITHUB and LINKEDIN, create hyperlinks instead of plain text
        """
        full_text = paragraph.text

        # Check if there are any placeholders
        has_placeholder = any(placeholder in full_text for placeholder in replacements.keys())
        if not has_placeholder:
            return

        # Build a mapping using paragraph.text to get correct positions
        # We need to account for BOTH runs and hyperlink elements
        char_map = []
        elements = []  # List of (element, type, text) tuples
        text_position = 0  # Current position in paragraph.text

        import sys
        for elem in paragraph._element:
            if elem.tag.endswith('}r'):  # Run element
                # Find corresponding run object
                for run in paragraph.runs:
                    if run._element == elem:
                        run_text = run.text
                        elements.append((run, 'run', run_text))
                        for char_offset in range(len(run_text)):
                            char_map.append((len(elements) - 1, 'run', char_offset))
                            text_position += 1
                        break
            elif elem.tag.endswith('}hyperlink'):  # Hyperlink element
                # Get the actual displayed text from paragraph.text
                # We'll extract just the visible text length by checking paragraph.text
                # Find how many chars this hyperlink contributes to paragraph.text
                hyperlink_runs = [r for r in elem.iter() if r.tag.endswith('}r')]
                hyperlink_text_parts = []
                for hr in hyperlink_runs:
                    for t in hr.iter():
                        if t.tag.endswith('}t') and t.text:
                            hyperlink_text_parts.append(t.text)
                            break  # Only take first text element per run

                hyperlink_display_text = ''.join(hyperlink_text_parts)
                elements.append((elem, 'hyperlink', hyperlink_display_text))
                for char_offset in range(len(hyperlink_display_text)):
                    char_map.append((len(elements) - 1, 'hyperlink', char_offset))
                    text_position += 1

        if not elements:
            return

        # Find all placeholders and their positions
        replacements_to_apply = []
        for placeholder, value in replacements.items():
            pos = full_text.find(placeholder)
            if pos != -1:
                replacements_to_apply.append((pos, placeholder, value))

        if not replacements_to_apply:
            return

        # Sort by position (process from start to end)
        replacements_to_apply.sort()

        # Process first placeholder
        pos, placeholder, value = replacements_to_apply[0]

        # Check if this is a hyperlink placeholder
        is_hyperlink = placeholder in ['{{GITHUB}}', '{{LINKEDIN}}']

        if is_hyperlink and cv_data:
            # Handle hyperlink specially
            if pos >= len(char_map):
                return

            start_elem_idx, start_elem_type, start_offset = char_map[pos]
            placeholder_end = pos + len(placeholder)

            if placeholder_end - 1 < len(char_map):
                end_elem_idx, end_elem_type, _ = char_map[placeholder_end - 1]
            else:
                end_elem_idx = start_elem_idx
                end_elem_type = start_elem_type

            # Get template run for formatting (find first run element)
            template_run = None
            for elem, elem_type, _ in elements:
                if elem_type == 'run':
                    template_run = elem
                    break

            if template_run is None:
                return

            # Get URL and display text
            if placeholder == '{{GITHUB}}':
                url = cv_data.get('personal_info', {}).get('github', '')
                display_text = 'Github'
            else:  # LINKEDIN
                url = cv_data.get('personal_info', {}).get('linkedin', '')
                display_text = 'Linkedin'

            # Build new content: elements before + hyperlink + elements after
            new_content = []

            # Elements before the placeholder
            for idx in range(start_elem_idx):
                elem, elem_type, elem_text = elements[idx]
                if elem_type == 'run':
                    if elem_text:
                        new_content.append(('text', elem_text, elem))
                elif elem_type == 'hyperlink':
                    new_content.append(('existing_hyperlink', elem))

            # Text before placeholder in the start element
            start_elem, start_elem_type, start_elem_text = elements[start_elem_idx]
            if start_elem_type == 'run':
                text_before = start_elem.text[:start_offset]
                if text_before:
                    new_content.append(('text', text_before, start_elem))

            # The new hyperlink
            new_content.append(('hyperlink', display_text, template_run, url))

            # Text after placeholder in the end element
            if end_elem_idx < len(elements):
                end_elem, end_elem_type, end_elem_text = elements[end_elem_idx]
                if end_elem_type == 'run':
                    # Calculate how many characters we've consumed
                    chars_before_end = sum(
                        len(elements[i][2])  # elements[i][2] is the text
                        for i in range(end_elem_idx)
                    )
                    end_offset = placeholder_end - chars_before_end
                    text_after = end_elem.text[end_offset:]
                    if text_after:
                        new_content.append(('text', text_after, end_elem))

            # Elements after the placeholder
            for idx in range(end_elem_idx + 1, len(elements)):
                elem, elem_type, elem_text = elements[idx]
                if elem_type == 'run':
                    if elem_text:
                        new_content.append(('text', elem_text, elem))
                elif elem_type == 'hyperlink':
                    new_content.append(('existing_hyperlink', elem))

            # Save paragraph properties before clearing
            para_alignment = paragraph.alignment
            para_style = paragraph.style

            # Save entire pPr (paragraph properties) element
            pPr = paragraph._element.pPr
            saved_pPr = None
            if pPr is not None:
                from copy import deepcopy
                saved_pPr = deepcopy(pPr)

            # Clear paragraph content (but keep pPr)
            for elem in list(paragraph._p):
                if not elem.tag.endswith('}pPr'):
                    paragraph._p.remove(elem)

            # If we had saved pPr, restore it
            if saved_pPr is not None:
                # Remove current pPr and replace with saved one
                current_pPr = paragraph._element.pPr
                if current_pPr is not None:
                    paragraph._p.remove(current_pPr)
                paragraph._p.insert(0, saved_pPr)

            # Rebuild paragraph with all elements
            for item in new_content:
                if item[0] == 'text':
                    _, text, source_run = item
                    new_run = paragraph.add_run(text)
                    new_run.font.name = source_run.font.name
                    new_run.font.size = source_run.font.size
                    new_run.font.bold = source_run.font.bold
                    new_run.font.italic = source_run.font.italic
                elif item[0] == 'hyperlink':
                    _, display_text, source_run, url = item
                    DOCXFiller._add_hyperlink(paragraph, display_text, url, source_run)
                elif item[0] == 'existing_hyperlink':
                    _, hyperlink_elem = item
                    # Re-append existing hyperlink element
                    paragraph._p.append(hyperlink_elem)

            return

        # Find which elements the placeholder spans (for non-hyperlink placeholders)
        placeholder_end = pos + len(placeholder)

        if pos >= len(char_map):
            return

        start_elem_idx, start_elem_type, start_offset = char_map[pos]
        end_elem_idx = start_elem_idx
        end_elem_type = start_elem_type

        if placeholder_end - 1 < len(char_map):
            end_elem_idx, end_elem_type, _ = char_map[placeholder_end - 1]

        # Get the template run for formatting
        start_elem, _, start_elem_text = elements[start_elem_idx]
        template_run = start_elem if start_elem_type == 'run' else None

        # If placeholder starts in hyperlink, find a run for template
        if template_run is None:
            for elem, elem_type, _ in elements:
                if elem_type == 'run':
                    template_run = elem
                    break

        if template_run is None:
            return

        # Build new content
        new_content = []

        # Elements before the placeholder
        for idx in range(start_elem_idx):
            elem, elem_type, elem_text = elements[idx]
            if elem_type == 'run':
                if elem_text:
                    new_content.append(('text', elem_text, elem))
            elif elem_type == 'hyperlink':
                new_content.append(('existing_hyperlink', elem))

        # Text before placeholder in the start element
        if start_elem_type == 'run':
            text_before = start_elem.text[:start_offset]
            if text_before:
                new_content.append(('text', text_before, start_elem))

        # The replacement value with template formatting
        new_content.append(('text', str(value), template_run))

        # Text after placeholder in the end element
        if end_elem_idx < len(elements):
            end_elem, end_elem_type, end_elem_text = elements[end_elem_idx]
            if end_elem_type == 'run':
                # Calculate how many characters we've consumed
                chars_before_end = sum(
                    len(elements[i][2])  # elements[i][2] is the text
                    for i in range(end_elem_idx)
                )
                end_offset = placeholder_end - chars_before_end
                text_after = end_elem.text[end_offset:]
                if text_after:
                    new_content.append(('text', text_after, end_elem))

        # Elements after the placeholder
        for idx in range(end_elem_idx + 1, len(elements)):
            elem, elem_type, elem_text = elements[idx]
            if elem_type == 'run':
                if elem_text:
                    new_content.append(('text', elem_text, elem))
            elif elem_type == 'hyperlink':
                new_content.append(('existing_hyperlink', elem))

        # Save paragraph properties before clearing
        para_alignment = paragraph.alignment
        para_style = paragraph.style

        # Save entire pPr (paragraph properties) element
        pPr = paragraph._element.pPr
        saved_pPr = None
        if pPr is not None:
            from copy import deepcopy
            saved_pPr = deepcopy(pPr)

        # Clear paragraph content (but keep pPr)
        for elem in list(paragraph._p):
            if not elem.tag.endswith('}pPr'):
                paragraph._p.remove(elem)

        # If we had saved pPr, restore it
        if saved_pPr is not None:
            # Remove current pPr and replace with saved one
            current_pPr = paragraph._element.pPr
            if current_pPr is not None:
                paragraph._p.remove(current_pPr)
            paragraph._p.insert(0, saved_pPr)

        # Rebuild paragraph with all elements
        for item in new_content:
            if item[0] == 'text':
                _, text, source_run = item
                new_run = paragraph.add_run(text)
                new_run.font.name = source_run.font.name
                new_run.font.size = source_run.font.size
                new_run.font.bold = source_run.font.bold
                new_run.font.italic = source_run.font.italic
                new_run.font.underline = source_run.font.underline
                if source_run.font.color and source_run.font.color.rgb:
                    new_run.font.color.rgb = source_run.font.color.rgb
            elif item[0] == 'existing_hyperlink':
                _, hyperlink_elem = item
                # Re-append existing hyperlink element
                paragraph._p.append(hyperlink_elem)


# Example usage and testing
if __name__ == "__main__":
    import sys

    if len(sys.argv) < 4:
        print("Usage: python src/docx_filler.py <template_path> <cv_yaml_path> <output_path>")
        print("Example: python src/docx_filler.py templates/template.docx outputs/customized_cv.yaml outputs/customized_cv.docx")
        sys.exit(1)

    template_path = sys.argv[1]
    cv_yaml_path = sys.argv[2]
    output_path = sys.argv[3]

    # Load CV data
    with open(cv_yaml_path, 'r') as f:
        cv_data = yaml.safe_load(f)

    print(f"Template: {template_path}")
    print(f"CV Data: {cv_yaml_path}")
    print(f"Output: {output_path}\n")

    # Fill template
    result = DOCXFiller.fill_template(template_path, cv_data, output_path)

    print(f"\n✓ Success! Resume generated at: {result}")
